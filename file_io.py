from datetime import datetime
import docx

def save_markdown(task_output):
  # Get today's date in the format YYYY-MM-DD
  today_date = datetime.now().strftime('%Y-%m-%d')

  # Set filenames with today's date
  filename_md = f"{today_date}.md"
  filename_docx = f"{today_date}.docx"

  result = task_output.result

  # Check if result is a method and call it if necessary
  if callable(result):
    result = result()

  # Convert result to string if it's not already
  if not isinstance(result, str):
    result = str(result)

  # Write Markdown file
  with open(filename_md, 'w', encoding='utf-8') as file:
    file.write(result)

  # Create a new Word document
  doc = docx.Document()

  # Handle dictionary output (specific format)
  if isinstance(result, dict):
    # Extract information from the dictionary
    source = result.get("Source")
    url = result.get("URL")
    published_date = result.get("Published Date")
    summary = result.get("Summary")

    # Create document content
    content = f"**Source:** {source}\n"
    if url:
      content += f"**URL:** {url}\n"
    if published_date:
      content += f"**Published Date:** {published_date}\n\n"
    content += summary

  # Handle non-dictionary output (existing logic with improvements)
  else:
    # Split the string result into lines
    lines = result.split('\n')

    # Iterate through each line and apply formatting
    content = ""
    for line in lines:
      if line.startswith('# '):
        # Use add_heading with level based on number of '#'
        level = len(line.split()[0]) - 1
        doc.add_heading(line[level + 1:], level)
      elif line.startswith('**') and line.endswith('**'):
        # Handle bold text
        bold_text = line[2:-2]
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(bold_text)
        run.bold = True
      else:
        doc.add_paragraph(line)

  # Save the document
  doc.save(filename_docx)
  print(f"Newsletter saved as {filename_docx}")
